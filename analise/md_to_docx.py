"""Converte o markdown do relatorio parcial (Paper 2) em .docx, embutindo as figuras
nos marcadores [[FIG:curva_ar]] e [[FIG:skill]]. Estilo limpo, sem exageros.
Uso: uv run --with python-docx python md_to_docx.py <entrada.md> <saida.docx>
"""
import re, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRATCH = Path(r"C:/Users/vinic/AppData/Local/Temp/claude/D--Artigo-JOH/bfcf1311-2eda-4171-b909-bd7d34f659a0/scratchpad")
FIGS = {
    "curva_ar": (SCRATCH / "curva_nse_lead_AR.png",
                 "Figura 1. NSE de teste por antecedencia (media de 5 sementes; faixa = desvio-padrao). "
                 "A persistencia (Lumped+AR) domina o curto prazo e cruza para baixo dos modelos com fisica no longo prazo."),
    "skill": (SCRATCH / "skill_forcante_58585000.png",
              "Figura 2. Skill da chuva prevista pelo GraphCast (GFS/IFS) na bacia 58585000, contra o telem observado: "
              "(a) correlacao por antecedencia; (b) vies de magnitude."),
}

md = Path(sys.argv[1]).read_text(encoding="utf-8")
out = sys.argv[2]

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(11)

def add_runs(par, text):
    """Interpreta **negrito** inline."""
    for i, seg in enumerate(re.split(r"(\*\*.+?\*\*)", text)):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            r = par.add_run(seg[2:-2]); r.bold = True
        else:
            par.add_run(seg)

def add_table(rows):
    rows = [r for r in rows if not re.match(r"^\s*\|[\s:|-]+\|\s*$", r)]  # tira separador ---
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            if j < len(t.rows[i].cells):
                cell = t.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val.replace("**", ""))
                if i == 0:
                    run.bold = True
    doc.add_paragraph()

def add_fig(key):
    path, cap = FIGS[key]
    if path.exists():
        doc.add_picture(str(path), width=Cm(16))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(cap); r.italic = True; r.font.size = Pt(9)
    else:
        doc.add_paragraph(f"[figura ausente: {path.name}]")

lines = md.splitlines()
i = 0
tbl_buf = []
while i < len(lines):
    ln = lines[i]
    if ln.strip().startswith("|"):
        tbl_buf.append(ln); i += 1; continue
    if tbl_buf:
        add_table(tbl_buf); tbl_buf = []
    m = re.match(r"\[\[FIG:(\w+)\]\]", ln.strip())
    if m:
        add_fig(m.group(1)); i += 1; continue
    if ln.startswith("### "):
        doc.add_heading(ln[4:].strip(), level=3)
    elif ln.startswith("## "):
        doc.add_heading(ln[3:].strip(), level=2)
    elif ln.startswith("# "):
        h = doc.add_heading(ln[2:].strip(), level=1)
    elif ln.strip() == "":
        pass
    elif re.match(r"^[-*] ", ln.strip()):
        p = doc.add_paragraph(style="List Bullet"); add_runs(p, ln.strip()[2:])
    else:
        p = doc.add_paragraph(); add_runs(p, ln)
    i += 1
if tbl_buf:
    add_table(tbl_buf)

doc.save(out)
print("docx salvo:", out)
